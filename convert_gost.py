"""
AURORA PNT — Конвертер Markdown → DOCX по ГОСТ 2.105-2019 / ГОСТ 7.32-2017.

Исправления:
  - Полный LaTeX→Unicode рендерер для инлайн ($...$) и блочных ($$...$$) формул
  - Устранение неразрывных пробелов из LaTeX-кода (~, backslash-comma, backslash-semi, quad)
  - Корректное ГОСТ-выравнивание всех элементов
  - Подписи таблиц ПЕРЕД, рисунков ПОД
  - Times New Roman 14, поля 30/15/20/20, красная строка 1.25, межстрочный 1.5
"""

import re, os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"c:\Users\SHIWA\Documents\LEOPath"
MD_IN    = os.path.join(BASE_DIR, "AURORA_PNT_Technical_Project.md")
DOCX_OUT = os.path.join(BASE_DIR, "AURORA_PNT_GOST.docx")
PDF_OUT  = os.path.join(BASE_DIR, "AURORA_PNT_GOST.pdf")

# ── ГОСТ-параметры ────────────────────────────────────────────────────────────
FONT_MAIN  = "Times New Roman"
FONT_MONO  = "Courier New"
SZ         = Pt(14)
SZ_SM      = Pt(12)
SZ_MONO    = Pt(11)
SZ_CAP     = Pt(12)   # подписи к рисункам/таблицам

BLACK = RGBColor(0, 0, 0)
GRAY  = RGBColor(0x55, 0x55, 0x55)
BLUE  = RGBColor(0x00, 0x00, 0x8B)
RED_S = RGBColor(0x7B, 0x00, 0x00)  # для кода

INDENT = Cm(1.25)   # красная строка

# ═══════════════════════════════════════════════════════════════════════════════
# LaTeX → Unicode рендерер
# ═══════════════════════════════════════════════════════════════════════════════

# Таблица замен: команда → символ
_GREEK = {
    "alpha":"α","beta":"β","gamma":"γ","delta":"δ","epsilon":"ε","zeta":"ζ",
    "eta":"η","theta":"θ","iota":"ι","kappa":"κ","lambda":"λ","mu":"μ",
    "nu":"ν","xi":"ξ","pi":"π","rho":"ρ","sigma":"σ","tau":"τ",
    "upsilon":"υ","phi":"φ","chi":"χ","psi":"ψ","omega":"ω",
    "Alpha":"Α","Beta":"Β","Gamma":"Γ","Delta":"Δ","Epsilon":"Ε",
    "Theta":"Θ","Lambda":"Λ","Mu":"Μ","Nu":"Ν","Xi":"Ξ","Pi":"Π",
    "Rho":"Ρ","Sigma":"Σ","Tau":"Τ","Phi":"Φ","Psi":"Ψ","Omega":"Ω",
    "varepsilon":"ε","varphi":"φ","varrho":"ρ","varsigma":"ς",
    "vartheta":"ϑ","varpi":"ϖ",
}

_MATH = {
    "approx":"≈","neq":"≠","leq":"≤","geq":"≥","ll":"≪","gg":"≫",
    "pm":"±","mp":"∓","times":"×","div":"÷","cdot":"·","circ":"∘",
    "infty":"∞","partial":"∂","nabla":"∇","forall":"∀","exists":"∃",
    "in":"∈","notin":"∉","subset":"⊂","supset":"⊃","cup":"∪","cap":"∩",
    "rightarrow":"→","leftarrow":"←","Rightarrow":"⇒","Leftarrow":"⇐",
    "leftrightarrow":"↔","uparrow":"↑","downarrow":"↓",
    "sum":"Σ","prod":"Π","int":"∫","oint":"∮",
    "sqrt":"√","pm":"±","ldots":"…","cdots":"⋯","vdots":"⋮",
    "oplus":"⊕","otimes":"⊗","perp":"⊥","parallel":"∥",
    "langle":"⟨","rangle":"⟩","lceil":"⌈","rceil":"⌉",
    "lfloor":"⌊","rfloor":"⌋",
    "geqslant":"≥","leqslant":"≤",
}

_TEXT_CMDS = {  # \text{...}, \mathrm{...} → contained text
    "text", "mathrm", "mathbf", "mathit", "mathsf", "mathtt",
    "boldsymbol", "vec", "hat", "tilde", "bar", "dot", "ddot",
    "overline", "underline", "widehat", "widetilde",
}

_SPACE_CMDS = {  # удалить, заменить обычным пробелом
    ",", ";", "!", ":", " ", "quad", "qquad",
}

_REMOVE_CMDS = {
    "left", "right", "big", "Big", "bigg", "Bigg",
    "displaystyle", "textstyle", "scriptstyle",
    "limits", "nolimits", "nonumber", "label",
}


def _strip_braces(s):
    """Убрать внешние фигурные скобки: {abc} → abc"""
    s = s.strip()
    if s.startswith("{") and s.endswith("}"):
        depth = 0
        for k, c in enumerate(s):
            if c == "{": depth += 1
            elif c == "}": depth -= 1
            if depth == 0 and k < len(s)-1:
                return s  # внешние скобки не единственные
        return s[1:-1]
    return s


def latex_to_text(src: str) -> str:
    """Преобразует LaTeX-строку в читаемый Unicode-текст для Word."""
    s = src.strip()

    # 1. Удалить окружения \begin{...}...\end{...}  (только маркеры, не содержимое)
    s = re.sub(r'\\begin\{[^}]+\}', '', s)
    s = re.sub(r'\\end\{[^}]+\}', '', s)

    # 2. \\ → перенос строки (в формулах — разделитель строк)
    s = s.replace("\\\\", "  ")

    # 3. Тильда LaTeX (неразрывный пробел) → обычный пробел
    s = s.replace("~", " ")

    # 4. \frac{числитель}{знаменатель} → (числитель/знаменатель)
    def replace_frac(m):
        num = latex_to_text(m.group(1))
        den = latex_to_text(m.group(2))
        return f"({num}/{den})"
    s = re.sub(r'\\frac\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}',
               replace_frac, s)

    # 5. \sqrt{...} → √(...)
    def replace_sqrt(m):
        inner = latex_to_text(m.group(1))
        return f"√({inner})"
    s = re.sub(r'\\sqrt\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', replace_sqrt, s)
    s = re.sub(r'\\sqrt\[.*?\]\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', replace_sqrt, s)

    # 6. Текстовые команды \text{...}, \mathrm{...} и т.д. → содержимое
    def replace_text_cmd(m):
        return latex_to_text(m.group(1))
    text_pat = r'\\(?:' + '|'.join(_TEXT_CMDS) + r')\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
    for _ in range(3):
        s = re.sub(text_pat, replace_text_cmd, s)

    # 7. Команды, которые надо просто удалить
    for cmd in _REMOVE_CMDS:
        s = re.sub(r'\\' + re.escape(cmd) + r'\b', '', s)

    # 8. Команды-пробелы → пробел
    for cmd in _SPACE_CMDS:
        s = re.sub(r'\\' + re.escape(cmd), ' ', s)

    # 9. Греческие и математические символы
    def replace_cmd(m):
        name = m.group(1)
        if name in _GREEK: return _GREEK[name]
        if name in _MATH:  return _MATH[name]
        return name  # оставить имя команды как есть (не потерять смысл)
    s = re.sub(r'\\([a-zA-Z]+)', replace_cmd, s)

    # 10. Нижние индексы:  x_{ab} → x_ab,   x_a → x_a
    s = re.sub(r'_\{([^{}]+)\}', lambda m: f"_{m.group(1)}", s)

    # 11. Верхние индексы: x^{ab} → x^ab,   x^a → x^a (short)
    #     Числовые степени: 10^{14} → 10¹⁴ (Unicode superscripts)
    _sup_map = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴',
                '5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹',
                '+':'⁺','-':'⁻','n':'ⁿ','i':'ⁱ'}
    def sup_repl(m):
        inner = m.group(1)
        result = "".join(_sup_map.get(c, c) for c in inner)
        return result
    s = re.sub(r'\^\{([^{}]+)\}', sup_repl, s)
    s = re.sub(r'\^([0-9nij])', lambda m: _sup_map.get(m.group(1), m.group(1)), s)

    # 12. Убрать оставшиеся одиночные фигурные скобки
    s = s.replace("{", "").replace("}", "")

    # 13. Убрать \, и другие одиночные обратные слэши перед пробелом или концом
    s = re.sub(r'\\([^a-zA-Z\s])', r'\1', s)
    s = re.sub(r'\\$', '', s)

    # 14. Схлопнуть множественные пробелы (но не переносы строк)
    s = re.sub(r'  +', ' ', s)

    return s.strip()


def process_inline_math(text: str) -> list:
    """
    Разбивает строку на части: [(is_math, content), ...]
    is_math=True  → формульный текст (моноширинный)
    is_math=False → обычный текст
    """
    # Паттерн: $...$  (не $$)
    # Сначала заменяем $$ чтобы не захватить
    parts = []
    # Временно защитим $$...$$ маркерами
    placeholder = "\x00BLOCK\x00"
    protected = re.sub(r'\$\$[^$]*\$\$', placeholder, text)

    segments = re.split(r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)', protected, flags=re.DOTALL)
    is_math = False
    for seg in segments:
        if seg == placeholder:
            parts.append((False, "[формула]"))
        else:
            parts.append((is_math, latex_to_text(seg) if is_math else seg))
        is_math = not is_math
    return parts


def clean_text(s: str) -> str:
    """Убирает все неразрывные и нулевые пробелы из обычного текста."""
    s = s.replace(" ", " ")   # nbsp
    s = s.replace(" ", " ")   # narrow nbsp
    s = s.replace(" ", " ")   # en space
    s = s.replace(" ", " ")   # em space
    s = s.replace("​", "")    # zero-width space
    s = s.replace("﻿", "")    # BOM
    return s


# ═══════════════════════════════════════════════════════════════════════════════
# Утилиты OXml
# ═══════════════════════════════════════════════════════════════════════════════

def set_spacing(p, line=1.5, before=0, after=0, first=None):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing  = line
    pf.space_before  = Pt(before)
    pf.space_after   = Pt(after)
    if first is not None:
        pf.first_line_indent = first

def fmt_run(run, size=None, bold=False, italic=False,
            color=None, font=None, mono=False):
    run.font.name  = (FONT_MONO if mono else (font or FONT_MAIN))
    run.font.size  = size or SZ
    run.bold       = bold
    run.italic     = italic
    run.font.color.rgb = color or BLACK

def add_page_num(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ftype, text in (("begin",""),("separate"," PAGE "),("end","")):
        r = fp.add_run(text if ftype == "separate" else "")
        r.font.name = FONT_MAIN
        r.font.size = SZ_SM
        if ftype != "separate":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), ftype)
            r._r.append(el)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            r._r.append(instr)

def cell_shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def table_borders(tbl_obj, color="AAAAAA"):
    tblPr = tbl_obj.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl_obj.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    brd = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        brd.append(b)
    tblPr.append(brd)

def set_col_widths(tbl_obj, widths_twips):
    grid = tbl_obj.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid"); tbl_obj.insert(1, grid)
    for old in grid.findall(qn("w:gridCol")):
        grid.remove(old)
    for w in widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w)))
        grid.append(gc)


# ═══════════════════════════════════════════════════════════════════════════════
# Добавление текста с инлайн-разметкой (**bold**, `code`, _italic_, $math$)
# ═══════════════════════════════════════════════════════════════════════════════

def add_inline(para, text: str, base_size=None, base_bold=False):
    """Добавляет текст в абзац, обрабатывая разметку и формулы."""
    text = clean_text(text)
    sz = base_size or SZ

    # Единый паттерн: **bold**, `code`, _italic_, $math$
    pattern = re.compile(
        r'(\*\*(?:[^*]|\*(?!\*))+\*\*'   # **bold**
        r'|`[^`]+`'                        # `code`
        r'|_[^_]+_'                        # _italic_
        r'|\$[^$\n]+\$'                    # $math$ (не $$)
        r')',
        re.DOTALL
    )
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = para.add_run(clean_text(part[2:-2]))
            fmt_run(r, size=sz, bold=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = para.add_run(clean_text(part[1:-1]))
            fmt_run(r, size=Pt(11), mono=True, color=RED_S)
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            r = para.add_run(clean_text(part[1:-1]))
            fmt_run(r, size=sz, italic=True)
        elif part.startswith("$") and part.endswith("$") and len(part) > 2:
            # Инлайн-формула
            formula_raw = part[1:-1]
            formula_text = latex_to_text(formula_raw)
            r = para.add_run(formula_text)
            fmt_run(r, size=sz, mono=True)
        else:
            # Обычный текст
            r = para.add_run(clean_text(part))
            fmt_run(r, size=sz, bold=base_bold)


# ═══════════════════════════════════════════════════════════════════════════════
# Счётчики ГОСТ
# ═══════════════════════════════════════════════════════════════════════════════

class Counters:
    def __init__(self):
        self.fig = 0; self.tbl = 0; self.eq = 0
        self.sec = [0, 0, 0]

    def nfig(self): self.fig += 1; return self.fig
    def ntbl(self): self.tbl += 1; return self.tbl
    def neq(self):  self.eq  += 1; return self.eq

    def h_num(self, level):
        if level == 1:
            self.sec[0] += 1; self.sec[1] = 0; self.sec[2] = 0
            return str(self.sec[0])
        elif level == 2:
            self.sec[1] += 1; self.sec[2] = 0
            return f"{self.sec[0]}.{self.sec[1]}"
        else:
            self.sec[2] += 1
            return f"{self.sec[0]}.{self.sec[1]}.{self.sec[2]}"


# ═══════════════════════════════════════════════════════════════════════════════
# Элементы документа ГОСТ
# ═══════════════════════════════════════════════════════════════════════════════

# Разделы нумеруются сквозно (1…64), как в исходнике и кросс-ссылках §N.
_UNNUMBERED = set()

def add_heading(doc, raw_text, level, cnt):
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', raw_text)  # убрать ссылки
    text = re.sub(r'^\d[\d.]*\s+', '', text.strip())           # убрать старую нумерацию

    unnumbered = text.strip().lower() in _UNNUMBERED
    num = "" if unnumbered else cnt.h_num(level)
    full = (f"{num}  {text}" if num else text)  # неразрывный пробел только в номере!
    # → нет, используем обычный пробел:
    full = (f"{num}  {text}" if num else text)

    # Стиль «Заголовок N» обязателен: по нему поле TOC собирает оглавление
    # (Ctrl+A → F9). Визуальное оформление ниже переопределяет стиль вручную.
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = Cm(0)

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, line=1.0, before=18, after=8)
        r = p.add_run(full.upper())
        fmt_run(r, size=SZ, bold=True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = INDENT
        set_spacing(p, line=1.0, before=14, after=6)
        r = p.add_run(full)
        fmt_run(r, size=SZ, bold=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = INDENT
        set_spacing(p, line=1.0, before=10, after=4)
        r = p.add_run(full)
        fmt_run(r, size=SZ, bold=True, italic=True)


def add_body_para(doc, text):
    """Обычный текстовый абзац с красной строкой, по ширине."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, line=1.5, before=0, after=0, first=INDENT)
    add_inline(p, text)
    return p


def add_quote_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    set_spacing(p, line=1.5, before=4, after=4)
    add_inline(p, text)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = GRAY
    return p


def add_list_item(doc, text, bullet=True, indent_level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    left = Cm(1.25 + indent_level * 0.6)
    p.paragraph_format.left_indent       = left
    p.paragraph_format.first_line_indent = Cm(-0.5)
    set_spacing(p, line=1.5, before=0, after=0)
    marker = "– " if bullet else ""
    r = p.add_run(marker)
    fmt_run(r, size=SZ)
    add_inline(p, text)
    return p


def add_numbered_item(doc, num_str, text, indent_level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    left = Cm(1.25 + indent_level * 0.6)
    p.paragraph_format.left_indent       = left
    p.paragraph_format.first_line_indent = Cm(-0.6)
    set_spacing(p, line=1.5, before=0, after=0)
    r = p.add_run(f"{num_str}) ")
    fmt_run(r, size=SZ)
    add_inline(p, text)
    return p


# Габариты «рамки» под рисунок (ГОСТ: ширина текста A4 при полях 30/15 = 165 мм)
FIG_MAX_W_CM = 16.0
FIG_MAX_H_CM = 21.0   # с запасом под подпись на A4 (поля 20/20 → 257 мм)


def _fig_width(img_path):
    """Ширина рисунка с сохранением пропорций, вписанная в рамку (см)."""
    try:
        from PIL import Image
        with Image.open(img_path) as im:
            w_px, h_px = im.size
        aspect = w_px / h_px if h_px else 1.4
        # по ширине рамки; если высота превысит лимит — ограничиваем по высоте
        w_cm = min(FIG_MAX_W_CM, FIG_MAX_H_CM * aspect)
        return Cm(w_cm)
    except Exception:
        return Cm(FIG_MAX_W_CM)


def add_figure(doc, img_path, alt, cnt):
    n = cnt.nfig()
    # Рисунок
    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.first_line_indent = Cm(0)
    set_spacing(pi, line=1.0, before=8, after=2)
    if os.path.exists(img_path):
        try:
            pi.add_run().add_picture(img_path, width=_fig_width(img_path))
        except Exception:
            r = pi.add_run(f"[Рисунок: {alt}]"); r.italic = True; fmt_run(r, color=GRAY)
    else:
        r = pi.add_run(f"[Файл не найден: {os.path.basename(img_path)}]")
        r.italic = True; fmt_run(r, color=GRAY)

    # Подпись (ГОСТ: под рисунком, по центру)
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    set_spacing(pc, line=1.0, before=2, after=12)
    r1 = pc.add_run(f"Рисунок {n}"); fmt_run(r1, size=SZ_CAP, bold=True)
    if alt:
        r2 = pc.add_run(f" — {clean_text(alt)}"); fmt_run(r2, size=SZ_CAP)


def add_table(doc, rows, caption, cnt):
    n = cnt.ntbl()

    # Подпись (ГОСТ: НАД таблицей, выравнивание влево)
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pc.paragraph_format.first_line_indent = Cm(0)
    set_spacing(pc, line=1.0, before=14, after=2)
    r1 = pc.add_run(f"Таблица {n}"); fmt_run(r1, size=SZ_CAP, bold=True)
    if caption:
        r2 = pc.add_run(f" — {clean_text(caption)}"); fmt_run(r2, size=SZ_CAP)

    # Фильтруем строки-разделители
    data = [r for r in rows if not re.match(r'^\|[\s:|\-]+\|$', r)]
    if not data:
        return

    n_cols = max(len(r.split("|")) - 2 for r in data)
    if n_cols <= 0:
        return

    # Ширина колонок
    total_twips = int(16.5 * 567)  # 16.5 cm в твипах (21 - 3 - 1.5 = 16.5)
    col_w = total_twips // n_cols

    tbl = doc.add_table(rows=len(data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_borders(tbl._tbl)
    set_col_widths(tbl._tbl, [col_w] * n_cols)

    for ri, row_str in enumerate(data):
        cells_txt = [c.strip() for c in row_str.split("|")[1:-1]]
        is_hdr = (ri == 0)
        for ci in range(n_cols):
            val = cells_txt[ci] if ci < len(cells_txt) else ""
            cell = tbl.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if is_hdr:
                cell_shade(cell, "D9E1F2")
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_hdr else WD_ALIGN_PARAGRAPH.LEFT
            cp.paragraph_format.first_line_indent = Cm(0)
            set_spacing(cp, line=1.0, before=2, after=2)
            add_inline(cp, val, base_size=SZ_SM, base_bold=is_hdr)

    # Пустой абзац после таблицы
    pa = doc.add_paragraph()
    pa.paragraph_format.first_line_indent = Cm(0)
    set_spacing(pa, line=1.0, before=0, after=6)


def add_formula_block(doc, latex_src, cnt):
    """Блочная формула: по центру, номер справа."""
    n = cnt.neq()
    sec = cnt.sec[0] or 0
    num_str = f"({sec}.{n})"

    rendered = latex_to_text(latex_src)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_spacing(p, line=1.2, before=6, after=6)

    r = p.add_run(rendered)
    r.font.name = FONT_MONO
    r.font.size = Pt(12)
    r.font.color.rgb = BLACK

    # Номер формулы в том же абзаце, справа через Tab (или просто отдельный маленький абзац)
    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pn.paragraph_format.first_line_indent = Cm(0)
    set_spacing(pn, line=1.0, before=0, after=6)
    rn = pn.add_run(num_str)
    fmt_run(rn, size=SZ)


def add_code_block(doc, lines_list):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.0)
    set_spacing(p, line=1.0, before=4, after=4)
    r = p.add_run("\n".join(lines_list))
    r.font.name = FONT_MONO; r.font.size = Pt(9)
    r.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


# ═══════════════════════════════════════════════════════════════════════════════
# Страницы: титул, реферат, оглавление
# ═══════════════════════════════════════════════════════════════════════════════

def _cp(doc, text, size=None, bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    set_spacing(p, line=1.0, before=before, after=after)
    r = p.add_run(clean_text(text))
    fmt_run(r, size=size or SZ, bold=bold, italic=italic)
    return p

def add_title_page(doc):
    _cp(doc, "РОССИЙСКАЯ ФЕДЕРАЦИЯ", bold=True, before=18)
    _cp(doc, "ООО «ШИВА НЕТВОРК» (ShiwaNetwork)", before=4)
    _cp(doc, "Основание: приказ ООО «ШИВА НЕТВОРК» № 11/2026", size=SZ_SM, before=4)
    _cp(doc, "")
    _cp(doc, "СОГЛАСОВАНО", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=18)
    _cp(doc, "Руководитель проекта", align=WD_ALIGN_PARAGRAPH.LEFT)
    _cp(doc, "________________  /________________/", align=WD_ALIGN_PARAGRAPH.LEFT)
    _cp(doc, '«____» ________________ 2026 г.', align=WD_ALIGN_PARAGRAPH.LEFT)
    _cp(doc, "")
    _cp(doc, "НАУЧНО-ИССЛЕДОВАТЕЛЬСКАЯ РАБОТА «СИЯНИЕ»", size=Pt(15), bold=True, before=22)
    _cp(doc, "")
    _cp(doc, "АВРОРА", size=Pt(22), bold=True, before=8)
    _cp(doc, "Система низкоорбитальной спутниковой навигации", size=Pt(16), bold=True)
    _cp(doc, "и высокоточной синхронизации", size=Pt(16), bold=True)
    _cp(doc, "")
    _cp(doc, "ТЕХНИЧЕСКИЙ ПРОЕКТ", size=Pt(18), bold=True, before=12)
    _cp(doc, "Шифр НИР: «Сияние».  Документ: СИЯНИЕ-ТП-001", bold=True, before=8)
    _cp(doc, "")
    _cp(doc, "Составлен в соответствии с ГОСТ 7.32-2017,", size=SZ_SM, italic=True, before=34)
    _cp(doc, "ГОСТ 2.105-2019, ГОСТ Р 7.0.97-2016", size=SZ_SM, italic=True)
    _cp(doc, "")
    _cp(doc, "Москва — 2026", bold=True, before=12)
    doc.add_page_break()

def add_abstract(doc, n_fig=0, n_tbl=0, n_src=0):
    _cp(doc, "РЕФЕРАТ", bold=True, before=0, after=12)

    # ГОСТ 7.32-2017 §6.3: реферат начинается со сведений об объёме отчёта.
    # Объём в страницах проставляется после финальной вёрстки в Word (F9).
    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ref.paragraph_format.first_line_indent = INDENT
    set_spacing(ref, line=1.5, before=0, after=10)
    r = ref.add_run(
        f"Отчёт о НИР: ____ с., {n_fig} рис., {n_tbl} табл., {n_src} источн."
    )
    fmt_run(r, bold=True)

    abstract = (
        "Настоящий технический проект научно-исследовательской работы «Сияние» "
        "(основание — приказ ООО «ШИВА НЕТВОРК» № 11/2026) описывает систему АВРОРА — "
        "низкоорбитальную спутниковую систему позиционирования, навигации и синхронизации "
        "(LEO PNT), функционирующую как высокоточное дополнение (усиление) к ГЛОНАСС "
        "и обеспечивающую независимую от GPS шкалу времени. Документ содержит аналитические "
        "модели и результаты численного моделирования по всем подсистемам: созвездие "
        "Walker 300/15/1 на орбите 1000 км, покрытие и геометрия (PDOP), навигационный "
        "сигнал L1/L5, синхронизация (девиация Аллана, перенос шкалы по ISL), межспутниковые "
        "каналы Ka-диапазона, ионосферная коррекция, орбитальное определение, RAIM/ARAIM, "
        "энергетика, масса, тепловой режим, анализ конъюнкций и план развёртывания по пяти фазам.\n\n"
        "Объект разработки: LEO-навигационная группировка из 300 КА.\n"
        "Расчётные показатели: точность позиционирования H-95% < 1 м (двухчастотный PPP, "
        "комбинированный режим с ГЛОНАСС; CEP ≈ 1 м), синхронизация UTC(SU) < 5 нс, "
        "доступность по территории РФ 100 %.\n"
        "Часовая архитектура: CSAC на всех 300 КА, space-Rb (Quantum-18) на 15 якорных КА, "
        "наземный эталон — группа активных водородных мазеров Ч1-1008 (UTC(SU)).\n\n"
        "Ключевые слова: LEO, PNT, GNSS, ГЛОНАСС, Walker, ISL, PDOP, UERE, RAIM, "
        "девиация Аллана, CSAC, АВРОРА, ГОСТ 7.32, спутниковая навигация."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT
    set_spacing(p, line=1.5, before=0, after=0)
    r = p.add_run(abstract)
    fmt_run(r)
    doc.add_page_break()

def add_toc(doc):
    _cp(doc, "СОДЕРЖАНИЕ", bold=True, before=0, after=14)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    set_spacing(p, line=1.5, before=0, after=0)
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    run._r.append(fb); run._r.append(instr); run._r.append(fs)
    ph = p.add_run("Оглавление формируется в Word: выделить всё (Ctrl+A) и нажать F9.")
    fmt_run(ph, size=Pt(11), italic=True, color=GRAY)
    fe_run = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    fe_run._r.append(fe)
    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    set_spacing(note, line=1.0, before=8, after=0)
    rn = note.add_run(
        "[Для обновления оглавления откройте в Word и нажмите Ctrl+A, затем F9]"
    )
    fmt_run(rn, size=Pt(9), italic=True, color=GRAY)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# Главный парсер MD → DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def parse(md_path, doc, cnt):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code = False;   code_buf = []
    in_block_math = False; math_buf = []
    in_table = False;  tbl_rows = []; tbl_cap = ""
    prev_was_heading = False

    def flush_table():
        nonlocal tbl_rows, tbl_cap, in_table
        add_table(doc, tbl_rows, tbl_cap, cnt)
        tbl_rows = []; tbl_cap = ""; in_table = False

    # Преамбула исходника (заголовок, подзаголовок, статус, ручное «Содержание»)
    # в DOCX не нужна: титульный лист, реферат и автособираемое оглавление
    # добавляются конвертером отдельно. Парсинг начинаем с первого
    # нумерованного раздела «## N. …».
    while i < len(lines) and not re.match(r'^##\s+\d+\.', lines[i].strip()):
        i += 1

    while i < len(lines):
        raw     = lines[i].rstrip("\n")
        stripped = raw.strip()

        # Конец таблицы при нетабличной строке
        if in_table and not stripped.startswith("|"):
            flush_table()

        # ── Блок кода ────────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                if code_buf:
                    add_code_block(doc, code_buf)
            i += 1; continue
        if in_code:
            code_buf.append(raw); i += 1; continue

        # ── Блочная формула $$ ────────────────────────────────────────────────
        if stripped == "$$":
            if not in_block_math:
                in_block_math = True; math_buf = []
            else:
                in_block_math = False
                add_formula_block(doc, "\n".join(math_buf), cnt)
                math_buf = []
            i += 1; continue
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            add_formula_block(doc, stripped[2:-2], cnt)
            i += 1; continue
        if in_block_math:
            math_buf.append(stripped); i += 1; continue

        # ── Разделитель ---  ──────────────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            i += 1; continue

        # ── Таблица ───────────────────────────────────────────────────────────
        if stripped.startswith("|"):
            in_table = True; tbl_rows.append(stripped)
            i += 1; continue

        # ── Рисунок ───────────────────────────────────────────────────────────
        img = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img:
            alt  = img.group(1)
            rel  = img.group(2).replace("/", os.sep)
            path = os.path.join(BASE_DIR, rel)
            add_figure(doc, path, alt, cnt)
            i += 1; continue

        # ── Заголовки ─────────────────────────────────────────────────────────
        hm = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if hm:
            # В исходнике разделы — «## N.» (главные) и «### N.M» (подразделы);
            # «#» — заголовок документа. Сдвигаем: ## → уровень 1, ### → уровень 2.
            lvl  = max(1, min(len(hm.group(1)) - 1, 3))
            text = hm.group(2)
            if text.strip().lower() == "содержание":
                i += 1; continue  # уже добавлено
            # Разрыв перед новым разделом 1-го уровня
            if lvl == 1 and cnt.sec[0] > 0:
                doc.add_page_break()
            add_heading(doc, text, lvl, cnt)
            i += 1; continue

        # ── Цитата >... ───────────────────────────────────────────────────────
        if stripped.startswith(">"):
            add_quote_para(doc, stripped.lstrip("> ").strip())
            i += 1; continue

        # ── Маркированный список ─────────────────────────────────────────────
        lm = re.match(r'^(\s*)([-*+])\s+(.*)', raw)
        if lm:
            lvl_i = len(lm.group(1)) // 2
            add_list_item(doc, lm.group(3), bullet=True, indent_level=lvl_i)
            i += 1; continue

        # ── Нумерованный список ───────────────────────────────────────────────
        nm = re.match(r'^(\s*)(\d+)[.)]\s+(.*)', raw)
        if nm:
            lvl_i = len(nm.group(1)) // 2
            add_numbered_item(doc, nm.group(2), nm.group(3), indent_level=lvl_i)
            i += 1; continue

        # ── Пустая строка ────────────────────────────────────────────────────
        if not stripped:
            i += 1; continue

        # ── Обычный абзац ────────────────────────────────────────────────────
        add_body_para(doc, stripped)
        i += 1

    if in_table:
        flush_table()


# ═══════════════════════════════════════════════════════════════════════════════
# Настройка документа
# ═══════════════════════════════════════════════════════════════════════════════

def setup_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(1.5)
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)

    ns = doc.styles["Normal"]
    ns.font.name = FONT_MAIN
    ns.font.size = SZ
    ns.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lvl in (1,2,3):
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name  = FONT_MAIN
        hs.font.size  = SZ
        hs.font.bold  = True
        hs.font.color.rgb = BLACK
        hs.paragraph_format.keep_with_next = True
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# Точка входа
# ═══════════════════════════════════════════════════════════════════════════════

def export_pdf(docx_path=DOCX_OUT, pdf_path=PDF_OUT):
    """Экспорт DOCX → PDF.

    Предпочтительно через Word COM (win32com): помимо PDF это обновляет все поля
    и оглавление — эквивалент Ctrl+A→F9, т.е. один из ручных шагов выпадает.
    Fallback — docx2pdf. На Linux/без Word/без pywin32 — мягкий пропуск
    (сборка docx не падает; важно для CI).
    """
    # 1) Word COM — обновляет оглавление и поля перед экспортом
    try:
        import pythoncom
        import win32com.client as win32
        pythoncom.CoInitialize()
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        try:
            doc = word.Documents.Open(docx_path)
            try:                                  # Ctrl+A → F9 (поля + оглавление)
                doc.Fields.Update()
                for toc in doc.TablesOfContents:
                    toc.Update()
            except Exception:
                pass
            doc.ExportAsFixedFormat(pdf_path, 17)  # 17 = wdExportFormatPDF
            doc.Close(False)
            kb = os.path.getsize(pdf_path) // 1024
            print(f"  PDF: {pdf_path} — {kb} КБ (Word COM; оглавление обновлено)")
            return True
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
    except Exception as e:
        print(f"  Word COM недоступен ({type(e).__name__}); пробую docx2pdf…")

    # 2) docx2pdf (тоже использует Word на Windows)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        kb = os.path.getsize(pdf_path) // 1024
        print(f"  PDF: {pdf_path} — {kb} КБ (docx2pdf)")
        return True
    except Exception as e:
        print(f"  PDF-экспорт пропущен ({type(e).__name__}): нет Word/pywin32/docx2pdf. "
              f"DOCX собран; PDF — вручную из Word (Сохранить как PDF).")
        return False


def main():
    print("AURORA PNT — ГОСТ 2.105-2019 / 7.32-2017")
    print("Инициализация...")
    doc = setup_doc()
    cnt = Counters()

    add_page_num(doc.sections[0])

    # Пред-сканирование исходника для сведений об объёме в реферате (ГОСТ 7.32).
    # Счёт рисунков/таблиц совпадает с фактической сборкой (проверено).
    md_text = open(MD_IN, encoding="utf-8").read()
    n_fig = len(re.findall(r'^!\[', md_text, re.M))
    n_tbl = len(re.findall(r'^\|[\s:|\-]+\|\s*$', md_text, re.M))
    msrc = re.search(r'^## \d+\.\s+Список литературы.*?(?=^## |\Z)',
                     md_text, re.M | re.S)
    n_src = len(re.findall(r'^\d+\.', msrc.group(0), re.M)) if msrc else 0

    print("  Титульный лист...")
    add_title_page(doc)
    print("  Реферат...")
    add_abstract(doc, n_fig=n_fig, n_tbl=n_tbl, n_src=n_src)
    print("  Оглавление...")
    add_toc(doc)

    print("  Парсинг основного текста...")
    parse(MD_IN, doc, cnt)

    print(f"\n  Статистика: рисунков={cnt.fig}, таблиц={cnt.tbl}, формул={cnt.eq}")
    print(f"  Сохранение: {DOCX_OUT}")
    doc.save(DOCX_OUT)
    kb = os.path.getsize(DOCX_OUT) // 1024
    print(f"  Готово! {kb} КБ")

    if "--no-pdf" not in sys.argv:
        print("  Экспорт PDF…")
        export_pdf()

if __name__ == "__main__":
    main()
