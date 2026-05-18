"""
Convert AURORA_PNT_Technical_Project.md to Word .docx
Uses python-docx with full formatting support.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import xml.etree.ElementTree as ET

BASE_DIR = r"c:\Users\SHIWA\Documents\LEOPath"
MD_FILE  = os.path.join(BASE_DIR, "AURORA_PNT_Technical_Project.md")
OUT_FILE = os.path.join(BASE_DIR, "AURORA_PNT_Technical_Project.docx")


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "B0B0B0")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_inline_formatting(paragraph, text):
    """Parse bold (**text**), inline code (`text`), and plain text."""
    # Split on bold/code markers
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def parse_md_to_docx(md_path, out_path):
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── Default style ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Heading styles
    h_colors = {
        1: RGBColor(0x0d, 0x3b, 0x66),  # dark blue
        2: RGBColor(0x15, 0x57, 0x99),
        3: RGBColor(0x1a, 0x73, 0xa8),
    }
    for lvl in (1, 2, 3):
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "Arial"
        hs.font.color.rgb = h_colors[lvl]
        hs.font.size = Pt(16 - lvl * 2)
        hs.font.bold = True

    # ── Read markdown ─────────────────────────────────────────────────────────
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []
    in_code = False
    code_lines = []

    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        # ── Code block ────────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                p.paragraph_format.left_indent = Cm(1.0)
                pf = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F4F4F4")
                pf.append(shd)
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── Separator / horizontal rule ────────────────────────────────────────
        if stripped in ("---", "***", "___") or re.match(r'^-{3,}$', stripped):
            doc.add_paragraph()
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if stripped.startswith("|"):
            table_rows.append(stripped)
            i += 1
            # Collect all table rows
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i].strip())
                i += 1
            # Parse table
            data_rows = [r for r in table_rows if not re.match(r'^\|[\s\-|]+\|$', r)]
            if data_rows:
                cols_n = max(len(r.split("|")) - 2 for r in data_rows)
                if cols_n > 0:
                    tbl = doc.add_table(rows=len(data_rows), cols=cols_n)
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                    set_table_borders(tbl)
                    for row_idx, row_str in enumerate(data_rows):
                        cells = [c.strip() for c in row_str.split("|")[1:-1]]
                        for col_idx in range(cols_n):
                            val = cells[col_idx] if col_idx < len(cells) else ""
                            cell = tbl.rows[row_idx].cells[col_idx]
                            p = cell.paragraphs[0]
                            add_inline_formatting(p, val)
                            p.style = doc.styles["Normal"]
                            p.runs[0].font.size = Pt(9) if p.runs else None
                            if row_idx == 0:
                                set_cell_bg(cell, "D6E4F0")
                                for run in p.runs:
                                    run.bold = True
                                    run.font.size = Pt(9)
                    doc.add_paragraph()
            table_rows = []
            continue

        # ── Image ─────────────────────────────────────────────────────────────
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt_text = img_match.group(1)
            rel_path = img_match.group(2)
            img_path = os.path.join(BASE_DIR, rel_path.replace("/", os.sep))
            if os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6.0))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Caption
                    cap = doc.add_paragraph(f"Рис. {alt_text}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                except Exception as e:
                    p = doc.add_paragraph(f"[Рисунок: {alt_text}]")
                    p.runs[0].italic = True
            else:
                p = doc.add_paragraph(f"[Рисунок (файл не найден): {rel_path}]")
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x00, 0x00)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if h_match:
            lvl = min(len(h_match.group(1)), 3)
            title_text = h_match.group(2)
            # Remove markdown link anchors from headings
            title_text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', title_text)
            if lvl == 1:
                p = doc.add_heading(level=1)
                p.clear()
                run = p.add_run(title_text)
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = h_colors[1]
                run.font.name = "Arial"
            else:
                p = doc.add_heading(title_text, level=lvl)
            i += 1
            continue

        # ── Block quote (> ...) ──────────────────────────────────────────────
        if stripped.startswith(">"):
            text = stripped.lstrip("> ")
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p.paragraph_format.left_indent = Cm(1.5)
            i += 1
            continue

        # ── Bullet / ordered list ─────────────────────────────────────────────
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', raw)
        if list_match:
            indent = len(list_match.group(1)) // 2
            text = list_match.group(3)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.0 + indent * 0.5)
            add_inline_formatting(p, text)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        add_inline_formatting(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")
    size_kb = os.path.getsize(out_path) // 1024
    print(f"Size:  {size_kb} KB")


if __name__ == "__main__":
    parse_md_to_docx(MD_FILE, OUT_FILE)
